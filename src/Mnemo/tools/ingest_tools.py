"""
Phase 2 — Ingestion de documents externes.
Gère l'extraction, le chunking et l'indexation de fichiers PDF, DOCX et TXT/MD.

Séparation volontaire de memory_tools.py :
  - memory_tools → mémoire personnelle (memory.md + chunks)
  - ingest_tools  → documents de référence (doc_chunks)

Les doc_chunks participent au retrieval hybride global via retrieve_all().

Formats supportés :
  - .pdf  → pypdf      (uv add pypdf)
  - .docx → python-docx (uv add python-docx)
  - .txt / .md → lecture directe, chunking par mots
"""
import hashlib
import sqlite3
import re

from pathlib import Path
from typing import Iterator

import numpy as np

# ── Dépendances optionnelles ─────────────────────────────────
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import docx as python_docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from Mnemo.context import get_data_dir as _get_data_dir
from Mnemo.tools.memory_tools import (
    EMBED_MODEL,
    CATEGORY_WEIGHTS,
    embed,
    cosine_similarity,
    compute_hash,
    sanitize_str,
)

def _db_path_default() -> Path:
    return _get_data_dir() / "memory.db"

# ── Config ────────────────────────────────────────────────────
CHUNK_SIZE     = 400   # Mots par chunk (PDF → texte brut)
CHUNK_OVERLAP  = 50    # Mots de chevauchement entre chunks
MIN_CHUNK_LEN  = 80    # Caractères minimum pour qu'un chunk soit indexé


# ══════════════════════════════════════════════════════════════
# Extraction PDF
# ══════════════════════════════════════════════════════════════

def extract_pdf_pages(path: Path) -> list[dict]:
    """
    Extrait le texte de chaque page d'un PDF.
    Retourne une liste de dicts {page: int, text: str}.
    Lève ImportError si pypdf n'est pas installé.
    Lève ValueError si le fichier n'est pas un PDF lisible.
    """
    if not HAS_PYPDF:
        raise ImportError(
            "pypdf n'est pas installé. Lance : uv add pypdf"
        )
    if not path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {path}")
    if path.suffix.lower() != ".pdf":
        raise ValueError(f"Format non supporté : {path.suffix} (attendu : .pdf)")

    reader = PdfReader(str(path))
    pages  = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = sanitize_str(text.strip())
        if text:
            pages.append({"page": i, "text": text})
    return pages


def get_pdf_page_count(path: Path) -> int:
    """Retourne le nombre de pages d'un PDF sans extraire le texte."""
    if not HAS_PYPDF:
        raise ImportError("pypdf n'est pas installé. Lance : uv add pypdf")
    return len(PdfReader(str(path)).pages)


# ══════════════════════════════════════════════════════════════
# Extraction DOCX
# ══════════════════════════════════════════════════════════════

def extract_docx_pages(path: Path) -> list[dict]:
    """
    Extrait le texte d'un fichier DOCX : paragraphes, headings et tableaux.
    Retourne une liste de dicts {page: int, text: str}.
    Le "numéro de page" est fictif (1 par bloc de ~500 mots) car DOCX
    n'a pas de notion native de page accessible sans Word.
    Lève ImportError si python-docx n'est pas installé.
    """
    if not HAS_DOCX:
        raise ImportError("python-docx n'est pas installé. Lance : uv add python-docx")
    if not path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Format non supporté : {path.suffix} (attendu : .docx)")

    doc    = python_docx.Document(str(path))
    blocks = []

    # ── Paragraphes et headings ───────────────────────────────
    for para in doc.paragraphs:
        text = sanitize_str(para.text.strip())
        if text:
            # Préfixe les headings pour préserver la structure sémantique
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                text  = f"{'#' * int(level) if level.isdigit() else '#'} {text}"
            blocks.append(text)

    # ── Tableaux ──────────────────────────────────────────────
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [sanitize_str(cell.text.strip()) for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(chr(10).join(rows))



    if not blocks:
        return []

    # Regroupe les blocs en "pages" fictives de ~500 mots
    pages       = []
    page_num    = 1
    word_count  = 0
    page_blocks = []
    WORDS_PER_FAKE_PAGE = 500

    for block in blocks:
        page_blocks.append(block)
        word_count += len(block.split())
        if word_count >= WORDS_PER_FAKE_PAGE:
            pages.append({"page": page_num, "text": chr(10).join(page_blocks)})


            page_num   += 1
            word_count  = 0
            page_blocks = []

    if page_blocks:
        pages.append({"page": page_num, "text": chr(10).join(page_blocks)})



    return pages


# ══════════════════════════════════════════════════════════════
# Extraction TXT / Markdown
# ══════════════════════════════════════════════════════════════

def extract_text_pages(path: Path) -> list[dict]:
    """
    Extrait le texte d'un fichier .txt ou .md.
    Pas de dépendance externe — lecture directe UTF-8.
    Retourne une liste de dicts {page: int, text: str}
    en découpant en blocs de ~500 mots (pages fictives).
    """
    if not path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {path}")
    if path.suffix.lower() not in (".txt", ".md"):
        raise ValueError(f"Format non supporté : {path.suffix} (attendu : .txt ou .md)")

    text = sanitize_str(path.read_text(encoding="utf-8", errors="ignore"))
    if not text.strip():
        return []

    # Découpe en blocs de ~500 mots (pages fictives)
    words               = text.split()
    WORDS_PER_FAKE_PAGE = 500
    pages               = []
    for i in range(0, len(words), WORDS_PER_FAKE_PAGE):
        chunk_words = words[i:i + WORDS_PER_FAKE_PAGE]
        pages.append({
            "page": len(pages) + 1,
            "text": " ".join(chunk_words),
        })
    return pages


# ══════════════════════════════════════════════════════════════
# Extraction Code source
# ══════════════════════════════════════════════════════════════

# Mapping extension → langage (pour le header des chunks et la détection)
CODE_EXTENSIONS: dict[str, str] = {
    ".py":   "python",
    ".js":   "javascript",
    ".ts":   "typescript",
    ".c":    "c",
    ".cpp":  "cpp",
    ".h":    "c_header",
    ".cs":   "csharp",
    ".java": "java",
    ".sh":   "bash",
    ".bash": "bash",
    ".ps1":  "powershell",
}

# Patterns regex de début de fonction/classe par langage
_FUNC_PATTERNS: dict[str, str] = {
    "javascript":  r"^(export\s+)?(async\s+)?function\s+\w+|^(export\s+)?const\s+\w+\s*=\s*(async\s+)?\(",
    "typescript":  r"^(export\s+)?(async\s+)?function\s+\w+|^(export\s+)?const\s+\w+\s*=\s*(async\s+)?\(",
    "c":           r"^\w[\w\s\*]+\s+\w+\s*\(",
    "cpp":         r"^\w[\w\s\*:<>]+\s+\w+\s*\(",
    "c_header":    r"^\w[\w\s\*:<>]+\s+\w+\s*\(",
    "csharp":      r"^\s*(public|private|protected|internal|static|override|virtual|async).*\s+\w+\s*\(",
    "java":        r"^\s*(public|private|protected|static|final|abstract|synchronized).*\s+\w+\s*\(",
    "bash":        r"^\w[\w\-]*\s*\(\s*\)|^function\s+\w+",
    "powershell":  r"^function\s+\w+",
}


def _split_by_ast(source: str, filename: str) -> list[str]:
    """
    Découpe un fichier Python en chunks par fonction/classe via l'AST.
    Retourne une liste de strings, chacune étant une fonction ou classe complète.
    Les toplevel statements hors fonctions/classes sont regroupés en un bloc 'module'.
    Note : pas de seuil MIN_CHUNK_LEN ici — une fonction de 2 lignes est un chunk valide.
    """
    import ast as _ast

    try:
        tree = _ast.parse(source)
    except SyntaxError:
        # Si l'AST échoue, on revient sur le découpage par lignes
        return _split_by_lines(source)

    lines   = source.splitlines()
    chunks  = []
    covered = set()  # lignes déjà dans un chunk

    # Seulement les noeuds de premier niveau (pas les méthodes imbriquées dans classes)
    for node in _ast.iter_child_nodes(tree):
        if isinstance(node, (_ast.FunctionDef, _ast.AsyncFunctionDef, _ast.ClassDef)):
            start = node.lineno - 1
            end   = node.end_lineno
            block = chr(10).join(lines[start:end]).strip()
            if block:
                chunks.append(block)
            covered.update(range(start, end))

    # Lignes hors fonctions/classes → bloc "module level"
    orphan_lines = [l for i, l in enumerate(lines) if i not in covered]
    orphan_text  = chr(10).join(orphan_lines).strip()
    if orphan_text:
        chunks.append(orphan_text)

    return chunks if chunks else _split_by_lines(source)


def _split_by_regex(source: str, language: str) -> list[str]:
    """
    Découpe un fichier source en chunks en détectant les débuts de
    fonctions/classes via regex. Universel pour les langages non-Python.
    Note : pas de seuil MIN_CHUNK_LEN ici — chaque fonction est un chunk valide.
    """
    import re as _re

    pattern = _FUNC_PATTERNS.get(language)
    if not pattern:
        return _split_by_lines(source)

    lines   = source.splitlines()
    chunks  = []
    current = []

    for line in lines:
        if current and _re.match(pattern, line):
            block = chr(10).join(current).strip()
            if block:
                chunks.append(block)
            current = [line]
        else:
            current.append(line)

    if current:
        block = chr(10).join(current).strip()
        if block:
            chunks.append(block)

    return chunks if chunks else _split_by_lines(source)


def _split_by_lines(source: str, lines_per_chunk: int = 60) -> list[str]:
    """Fallback : découpe en blocs de N lignes."""
    lines  = source.splitlines()
    chunks = []
    for i in range(0, len(lines), lines_per_chunk):
        block = chr(10).join(lines[i:i + lines_per_chunk]).strip()
        if len(block) >= MIN_CHUNK_LEN:
            chunks.append(block)
    return chunks


def extract_code_pages(path: Path) -> list[dict]:
    """
    Extrait les chunks d'un fichier source.
    - Python : découpage AST (fonctions + classes)
    - Autres : découpage regex sur les patterns de fonctions/classes
    Retourne une liste de dicts {page: int, text: str}.
    Le "numéro de page" correspond à l'index du chunk dans le fichier.
    """
    if not path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {path}")

    ext      = path.suffix.lower()
    language = CODE_EXTENSIONS.get(ext)
    if not language:
        raise ValueError(f"Extension non supportée : {ext}")

    source = sanitize_str(path.read_text(encoding="utf-8", errors="ignore"))
    if not source.strip():
        return []

    if language == "python":
        raw_chunks = _split_by_ast(source, path.name)
    else:
        raw_chunks = _split_by_regex(source, language)

    # Chaque chunk devient une "page" avec un header indiquant le langage et le fichier
    pages = []
    for i, chunk in enumerate(raw_chunks):
        header = f"# [{language}] {path.name}"
        pages.append({
            "page": i + 1,
            "text": f"{header}\n{chunk}",

        })
    return pages


# ══════════════════════════════════════════════════════════════
# Chunking
# ══════════════════════════════════════════════════════════════

def clean_text(text: str) -> str:
    """
    Nettoie le texte extrait d'un PDF :
    - Supprime les lignes vides multiples
    - Normalise les espaces
    - Supprime les artefacts courants (numéros de page seuls, tirets de césure)
    """
    # Joins les mots coupés en fin de ligne (trait d'union + saut de ligne)
    text = re.sub(r"-\n", "", text)
    # Normalise les sauts de ligne multiples
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Normalise les espaces multiples
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> Iterator[str]:
    """
    Découpe un texte en chunks de `chunk_size` mots avec `overlap` mots
    de chevauchement entre chunks consécutifs.
    Préserve les frontières de phrases quand c'est possible.
    """
    words = text.split()
    if not words:
        return

    start = 0
    while start < len(words):
        end   = min(start + chunk_size, len(words))
        chunk = " ".join(words[start:end])
        if len(chunk) >= MIN_CHUNK_LEN:
            yield chunk
        start += chunk_size - overlap
        if start >= len(words):
            break


def chunk_pages(pages: list[dict]) -> list[dict]:
    """
    Découpe les pages extraites en chunks indexables.
    Retourne une liste de dicts :
    {page: int, chunk_index: int, content: str}
    """
    chunks      = []
    chunk_index = 0
    for page_data in pages:
        text = clean_text(page_data["text"])
        for chunk in chunk_text(text):
            chunks.append({
                "page":        page_data["page"],
                "chunk_index": chunk_index,
                "content":     chunk,
            })
            chunk_index += 1
    return chunks


# ══════════════════════════════════════════════════════════════
# Gestion des documents en DB
# ══════════════════════════════════════════════════════════════

def file_hash(path: Path) -> str:
    """MD5 du contenu du fichier — sert d'ID unique pour le document."""
    return hashlib.md5(path.read_bytes()).hexdigest()


def is_already_ingested(db: sqlite3.Connection, doc_id: str) -> bool:
    """Retourne True si ce document (même hash) a déjà été ingéré."""
    return db.execute(
        "SELECT 1 FROM documents WHERE id = ?", (doc_id,)
    ).fetchone() is not None


def register_document(
    db: sqlite3.Connection,
    doc_id: str,
    path: Path,
    page_count: int,
    chunk_count: int,
) -> None:
    """Enregistre le document dans la table documents."""
    db.execute("""
        INSERT INTO documents (id, filename, path, mime_type, page_count, chunk_count)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (doc_id, path.name, str(path.resolve()), "application/pdf", page_count, chunk_count))


def upsert_doc_chunk(
    db: sqlite3.Connection,
    doc_id: str,
    chunk: dict,
    filename: str,
) -> None:
    """
    Insère un chunk de document dans doc_chunks + doc_embeddings + doc_chunks_fts.
    Le chunk_id est le MD5 de (doc_id + page + chunk_index + contenu) —
    garantit l'unicité même si deux passages du document ont le même texte.
    """
    # BUG CORRIGÉ : hash sur contenu seul → collision si même texte apparaît
    # deux fois dans le document (ex: headers répétés, formules identiques)
    chunk_id = compute_hash(f"{doc_id}:{chunk['page']}:{chunk['chunk_index']}:{chunk['content']}")

    # Idempotence — on ne réinsère pas un chunk identique
    if db.execute("SELECT 1 FROM doc_chunks WHERE id = ?", (chunk_id,)).fetchone():
        return

    db.execute("""
        INSERT INTO doc_chunks
            (id, doc_id, page, chunk_index, content, importance_weight, category)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (
        chunk_id, doc_id,
        chunk["page"], chunk["chunk_index"],
        chunk["content"],
        CATEGORY_WEIGHTS.get("connaissance", 1.0),
        "connaissance",
    ))

    # Embedding
    vector = embed(chunk["content"])
    db.execute("""
        INSERT OR REPLACE INTO doc_embeddings (chunk_id, model, vector, dim)
        VALUES (?, ?, ?, ?)
    """, (chunk_id, EMBED_MODEL, vector.tobytes(), len(vector)))

    # Index FTS5
    db.execute("""
        INSERT INTO doc_chunks_fts (chunk_id, content, filename)
        VALUES (?, ?, ?)
    """, (chunk_id, chunk["content"], filename))


# ══════════════════════════════════════════════════════════════
# Pipeline principal
# ══════════════════════════════════════════════════════════════

def ingest_pdf(path: Path, db_path: Path | None = None) -> dict:
    """
    Pipeline complet d'ingestion d'un PDF :
    1. Vérifie que le fichier n'est pas déjà ingéré (par hash)
    2. Extrait le texte page par page
    3. Découpe en chunks
    4. Génère les embeddings et indexe en DB

    Retourne un dict de résumé :
    {
        "status":      "ingested" | "already_ingested" | "empty",
        "doc_id":      str,
        "filename":    str,
        "pages":       int,
        "chunks":      int,
    }
    """
    if db_path is None:
        db_path = _db_path_default()
    db     = sqlite3.connect(str(db_path))
    doc_id = file_hash(path)

    # ── Déjà ingéré ? ─────────────────────────────────────────
    if is_already_ingested(db, doc_id):
        db.close()
        return {
            "status":   "already_ingested",
            "doc_id":   doc_id,
            "filename": path.name,
            "pages":    0,
            "chunks":   0,
        }

    # ── Extraction ────────────────────────────────────────────
    pages      = extract_pdf_pages(path)
    page_count = get_pdf_page_count(path)

    if not pages:
        db.close()
        return {
            "status":   "empty",
            "doc_id":   doc_id,
            "filename": path.name,
            "pages":    page_count,
            "chunks":   0,
        }

    # ── Chunking ──────────────────────────────────────────────
    chunks = chunk_pages(pages)

    # ── Indexation ────────────────────────────────────────────
    total = len(chunks)
    for i, chunk in enumerate(chunks, start=1):
        upsert_doc_chunk(db, doc_id, chunk, path.name)
        # Progression toutes les 10 chunks ou sur le dernier
        if i % 10 == 0 or i == total:
            pct = int(i / total * 100)
            bar = "█" * (pct // 5) + "░" * (20 - pct // 5)
            print(f"\r   [{bar}] {pct:3d}%  chunk {i}/{total}", end="", flush=True)

    print()  # Saut de ligne final

    register_document(db, doc_id, path, page_count, len(chunks))
    db.commit()
    db.close()

    return {
        "status":   "ingested",
        "doc_id":   doc_id,
        "filename": path.name,
        "pages":    page_count,
        "chunks":   len(chunks),
    }


def _ingest_pages(
    path: Path,
    pages: list[dict],
    page_count: int,
    mime_type: str,
    db_path: Path | None = None,
) -> dict:
    """
    Pipeline d'indexation commun à tous les formats.
    Appelé par ingest_pdf, ingest_docx, ingest_text après extraction.
    """
    if db_path is None:
        db_path = _db_path_default()
    db     = sqlite3.connect(str(db_path))
    doc_id = file_hash(path)

    if is_already_ingested(db, doc_id):
        db.close()
        return {"status": "already_ingested", "doc_id": doc_id,
                "filename": path.name, "pages": 0, "chunks": 0}

    if not pages:
        db.close()
        return {"status": "empty", "doc_id": doc_id,
                "filename": path.name, "pages": page_count, "chunks": 0}

    chunks = chunk_pages(pages)
    total  = len(chunks)

    for i, chunk in enumerate(chunks, start=1):
        upsert_doc_chunk(db, doc_id, chunk, path.name)
        if i % 10 == 0 or i == total:
            pct = int(i / total * 100)
            bar = "█" * (pct // 5) + "░" * (20 - pct // 5)
            print(f"\r   [{bar}] {pct:3d}%  chunk {i}/{total}", end="", flush=True)

    print()

    # Surcharge mime_type dans documents pour DOCX/TXT
    db.execute("""
        INSERT INTO documents (id, filename, path, mime_type, page_count, chunk_count)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (doc_id, path.name, str(path.resolve()), mime_type, page_count, total))
    db.commit()
    db.close()

    return {"status": "ingested", "doc_id": doc_id,
            "filename": path.name, "pages": page_count, "chunks": total}


def ingest_docx(path: Path, db_path: Path | None = None) -> dict:
    """Ingère un fichier DOCX dans la base de connaissances."""
    pages      = extract_docx_pages(path)
    page_count = len(pages)  # pages fictives
    return _ingest_pages(path, pages, page_count, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", db_path)


def ingest_text(path: Path, db_path: Path | None = None) -> dict:
    """Ingère un fichier TXT ou Markdown dans la base de connaissances."""
    pages      = extract_text_pages(path)
    page_count = len(pages)  # pages fictives
    mime       = "text/markdown" if path.suffix.lower() == ".md" else "text/plain"
    return _ingest_pages(path, pages, page_count, mime, db_path)


def ingest_code(path: Path, db_path: Path | None = None) -> dict:
    """Ingère un fichier source dans la base de connaissances."""
    pages      = extract_code_pages(path)
    page_count = len(pages)
    language   = CODE_EXTENSIONS.get(path.suffix.lower(), "code")
    mime       = f"text/x-{language}"
    return _ingest_pages(path, pages, page_count, mime, db_path)


def ingest_text_block(
    text:     str,
    title:    str | None = None,
    source:   str = "text_block",
    db_path:  Path | None = None,
) -> dict:
    """
    Ingère un bloc de texte brut (sans fichier physique) dans doc_chunks.

    Utilisé par NoteWriterCrew pour les contenus classifiés Bucket B.
    Le doc_id est le MD5 du texte — garantit l'idempotence si le même
    contenu est soumis deux fois.

    Args:
        text   : contenu brut à ingérer.
        title  : titre optionnel (affiché dans le retrieval). Si absent,
                 généré depuis les premiers mots du texte.
        source : identifiant du type de source (stocké dans la colonne path
                 de la table documents). Défaut : "text_block".
        db_path: chemin vers memory.db. Défaut : DATA_DIR/memory.db.

    Returns:
        dict { status, doc_id, filename, pages, chunks }
        status ∈ "ingested" | "already_ingested" | "empty"
    """
    import hashlib as _hashlib

    if db_path is None:
        db_path = _db_path_default()

    cleaned = sanitize_str(text)
    if not cleaned.strip():
        return {"status": "empty", "doc_id": "", "filename": title or "", "pages": 0, "chunks": 0}

    doc_id = _hashlib.md5(cleaned.encode()).hexdigest()

    if not title:
        words = cleaned.split()[:8]
        title = " ".join(words) + ("…" if len(cleaned.split()) > 8 else "")

    db = sqlite3.connect(str(db_path))

    if is_already_ingested(db, doc_id):
        db.close()
        return {"status": "already_ingested", "doc_id": doc_id, "filename": title, "pages": 0, "chunks": 0}

    # Découpe en pages fictives de 500 mots (même logique que extract_text_pages)
    WORDS_PER_PAGE = 500
    words = cleaned.split()
    pages = []
    for i in range(0, len(words), WORDS_PER_PAGE):
        pages.append({
            "page": len(pages) + 1,
            "text": " ".join(words[i:i + WORDS_PER_PAGE]),
        })

    if not pages:
        db.close()
        return {"status": "empty", "doc_id": doc_id, "filename": title, "pages": 0, "chunks": 0}

    chunks = chunk_pages(pages)
    total  = len(chunks)

    for chunk in chunks:
        upsert_doc_chunk(db, doc_id, chunk, title)

    db.execute("""
        INSERT INTO documents (id, filename, path, mime_type, page_count, chunk_count)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (doc_id, title, source, "text/plain", len(pages), total))
    db.commit()
    db.close()

    return {"status": "ingested", "doc_id": doc_id, "filename": title, "pages": len(pages), "chunks": total}


def ingest_file(path: Path, db_path: Path | None = None) -> dict:
    """
    Dispatcher universel — détecte le format et appelle le bon ingester.
    Formats supportés : .pdf, .docx, .txt, .md, .py, .js, .ts,
                        .c, .cpp, .h, .cs, .java, .sh, .bash, .ps1
    """
    ext = path.suffix.lower()
    if ext == ".pdf":
        return ingest_pdf(path, db_path)
    elif ext == ".docx":
        return ingest_docx(path, db_path)
    elif ext in (".txt", ".md"):
        return ingest_text(path, db_path)
    elif ext in CODE_EXTENSIONS:
        return ingest_code(path, db_path)
    else:
        raise ValueError(
            f"Format non supporté : {ext}\n"
            f"Formats acceptés : .pdf, .docx, .txt, .md, "
            f".py, .js, .ts, .c, .cpp, .h, .cs, .java, .sh, .bash, .ps1"
        )




def delete_document(doc_id: str, db_path: Path | None = None) -> bool:
    """
    Supprime un document et tous ses chunks (cascade) de la base.
    Nettoie aussi doc_chunks_fts (table virtuelle sans FK constraint).
    Retourne True si le document existait, False s'il était introuvable.
    """
    if db_path is None:
        db_path = _db_path_default()
    db = sqlite3.connect(str(db_path))
    row = db.execute("SELECT id FROM documents WHERE id = ?", (doc_id,)).fetchone()
    if not row:
        db.close()
        return False
    # Récupère les chunk_ids avant suppression pour nettoyer FTS
    chunk_ids = [r[0] for r in db.execute(
        "SELECT id FROM doc_chunks WHERE doc_id = ?", (doc_id,)
    ).fetchall()]
    for cid in chunk_ids:
        db.execute("DELETE FROM doc_chunks_fts WHERE chunk_id = ?", (cid,))
    # ON DELETE CASCADE supprime doc_chunks + doc_embeddings
    db.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
    db.commit()
    db.close()
    return True


# ══════════════════════════════════════════════════════════════
# Retrieval dans les documents (utilisé par retrieve_all)
# ══════════════════════════════════════════════════════════════

def search_docs_keyword(db: sqlite3.Connection, query: str, top_k: int = 10) -> list[dict]:
    """Recherche FTS5 dans les doc_chunks."""
    if not query or not query.strip():
        return []
    from Mnemo.tools.memory_tools import _sanitize_fts_query
    fts_query = _sanitize_fts_query(query)
    if not fts_query:
        return []
    rows = db.execute("""
        SELECT dc.id, dc.content, dc.page, dc.chunk_index,
               d.filename, bm25(doc_chunks_fts) as score,
               dc.importance_weight, dc.category
        FROM doc_chunks_fts
        JOIN doc_chunks dc  ON doc_chunks_fts.chunk_id = dc.id
        JOIN documents  d   ON dc.doc_id = d.id
        WHERE doc_chunks_fts MATCH ?
        ORDER BY score
        LIMIT ?
    """, (fts_query, top_k)).fetchall()
    return [{
        "id":               r[0],
        "content":          r[1],
        "page":             r[2],
        "chunk_index":      r[3],
        "source":           r[4],   # filename — affiché dans le prompt
        "score_fts":        abs(r[5]),
        "importance_weight": r[6],
        "category":         r[7],
        "type":             "document",
    } for r in rows]


def search_docs_vector(db: sqlite3.Connection, query: str, top_k: int = 10) -> list[dict]:
    """Recherche vectorielle dans les doc_chunks."""
    query_vec = embed(query, prefix="search_query")
    rows = db.execute("""
        SELECT dc.id, dc.content, dc.page, dc.chunk_index,
               d.filename, de.vector,
               dc.importance_weight, dc.category
        FROM doc_embeddings de
        JOIN doc_chunks  dc ON de.chunk_id = dc.id
        JOIN documents    d ON dc.doc_id   = d.id
    """).fetchall()
    results = []
    for r in rows:
        vec = np.frombuffer(r[5], dtype=np.float32)
        results.append({
            "id":               r[0],
            "content":          r[1],
            "page":             r[2],
            "chunk_index":      r[3],
            "source":           r[4],
            "score_vector":     cosine_similarity(query_vec, vec),
            "importance_weight": r[6],
            "category":         r[7],
            "type":             "document",
        })
    results.sort(key=lambda x: x["score_vector"], reverse=True)
    return results[:top_k]


# ══════════════════════════════════════════════════════════════
# Utilitaires CLI
# ══════════════════════════════════════════════════════════════

def list_ingested_documents(db_path: Path | None = None) -> list[dict]:
    """Retourne la liste des documents ingérés avec leurs métadonnées."""
    if db_path is None:
        db_path = _db_path_default()
    db   = sqlite3.connect(str(db_path))
    rows = db.execute("""
        SELECT id, filename, page_count, chunk_count, ingested_at
        FROM documents
        ORDER BY ingested_at DESC
    """).fetchall()
    db.close()
    return [
        {
            "doc_id":      r[0],
            "filename":    r[1],
            "pages":       r[2],
            "chunks":      r[3],
            "ingested_at": r[4],
        }
        for r in rows
    ]